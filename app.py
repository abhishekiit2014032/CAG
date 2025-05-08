import os
import datetime
from dotenv import load_dotenv
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
import streamlit as st

# Load environment variables from .env file
load_dotenv()

# ========== STEP 1: Setup API Key ==========
API_KEY = os.environ.get("GEMINI_API_KEY")
if not API_KEY:
    st.error("Please set the GEMINI_API_KEY environment variable.")
    st.stop()

genai.configure(api_key=API_KEY)
MODEL_RESOURCE_NAME = "models/gemini-1.5-pro-001"

# ========== STEP 2: Document Extraction ==========

def extract_text_from_pdf(pdf_file):
    """Extracts text from all pages of a PDF file."""
    try:
        reader = PdfReader(pdf_file)
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(f"[Page {i+1}]\n{text}")
        
        st.info(f"Successfully extracted {len(text_parts)} pages from PDF: {pdf_file.name}")
        return "\n\n".join(text_parts)
    except Exception as e:
        st.error(f"Error reading PDF {pdf_file.name}: {e}")
        return ""

def extract_text_from_docx(docx_file):
    """Extracts text from all paragraphs of a DOCX file."""
    try:
        doc = Document(docx_file)
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        st.info(f"Successfully extracted {len(text_parts)} text elements from DOCX: {docx_file.name}")
        return "\n\n".join(text_parts)
    except Exception as e:
        st.error(f"Error reading DOCX {docx_file.name}: {e}")
        return ""

# ========== STEP 3: Create Cached Context ==========

def process_document_content(pdf_file=None, docx_file=None):
    """
    Extracts text from specified PDF and/or DOCX files and returns the combined content.
    """
    combined_content = []

    if pdf_file is not None:
        st.info(f"Extracting PDF content from: {pdf_file.name}")
        pdf_text = extract_text_from_pdf(pdf_file)
        if pdf_text.strip():
            combined_content.append(pdf_text)

    if docx_file is not None:
        st.info(f"Extracting DOCX content from: {docx_file.name}")
        docx_text = extract_text_from_docx(docx_file)
        if docx_text.strip():
            combined_content.append(docx_text)

    final_content_string = "\n\n--- End of Document Section ---\n\n".join(combined_content)

    # Check if any content was extracted
    if not final_content_string.strip():
        st.warning("No content extracted from files.")
        return None

    st.info(f"Total extracted content length: {len(final_content_string)} characters.")
    return final_content_string

# ========== STEP 4: Chat Interface with Streamlit ==========

def chatbot_interface(document_content, language="English"):
    """
    Initializes a GenerativeModel and provides a chat interface using Streamlit.
    The document content is used as context for the model.
    Supports both English and Marathi languages.
    """
    if not document_content:
        error_msg = "दस्तऐवज सामग्री उपलब्ध नसल्यामुळे चॅटबॉट सुरू करू शकत नाही." if language == "मराठी" else "Cannot start chatbot as document content is not available."
        st.error(error_msg)
        return

    try:
        # Use the most capable model available
        model = genai.GenerativeModel(MODEL_RESOURCE_NAME)
        
        # Display ready message in the selected language
        ready_msg = "✅ चॅटबॉट तयार आहे. आपल्या प्रश्नांची उत्तरे अपलोड केलेल्या दस्तऐवजांच्या आधारे दिली जातील." if language == "मराठी" else "✅ Chatbot is ready. Your questions will be answered based on the uploaded documents."
        st.info(ready_msg)

        # Display document content length for debugging
        length_msg = f"दस्तऐवज सामग्री लांबी: {len(document_content)} वर्ण" if language == "मराठी" else f"Document content length: {len(document_content)} characters"
        st.info(length_msg)
        
        # Initialize chat session with document content
        if "chat" not in st.session_state:
            # Add system instructions as a user message in the appropriate language
            if language == "मराठी":
                system_message = (
                    "तुम्ही एक मदतगार सहाय्यक आहात. "
                    "तुमचे प्राथमिक कार्य प्रदान केलेल्या दस्तऐवज संदर्भावर आधारित वापरकर्त्याच्या प्रश्नांची उत्तरे देणे आहे. "
                    "जर प्रश्नाचे उत्तर प्रदान केलेल्या दस्तऐवजांमध्ये आढळले नाही, तर तुम्ही स्पष्टपणे सांगा: 'ही माहिती प्रदान केलेल्या दस्तऐवजांमध्ये उपलब्ध नाही.' "
                    "जर माहिती दस्तऐवजांमध्ये नसेल तर सामान्य ज्ञानातून उत्तर देण्याचा प्रयत्न करू नका. "
                    "तुम्ही मराठी भाषेत उत्तर द्या."
                )
                
                document_message = (
                    "दस्तऐवज सामग्री सुरू\n\n"
                    f"{document_content}\n\n"
                    "दस्तऐवज सामग्री समाप्त\n\n"
                    "कृपया माझ्या प्रश्नांची उत्तरे देण्यासाठी फक्त दस्तऐवज सामग्री सुरू आणि दस्तऐवज सामग्री समाप्त यांच्या दरम्यान असलेली माहिती वापरा. "
                    "जर माहिती दस्तऐवजात नसेल तर तसे सांगा."
                )
                
                model_response1 = "मी माझी भूमिका समजून घेतली आहे. मी फक्त दस्तऐवज सामग्रीवर आधारित प्रश्नांची उत्तरे देईन."
                model_response2 = "मला दस्तऐवज सामग्री मिळाली आहे आणि मी तुमच्या प्रश्नांची उत्तरे देण्यासाठी त्याचा वापर करेन."
            else:
                system_message = (
                    "You are a helpful assistant. "
                    "Your primary function is to answer user questions strictly based on the "
                    "information contained in the provided document context. "
                    "If the answer to a question cannot be found within the provided documents, "
                    "you MUST explicitly state: 'The information is not available in the provided documents.' "
                    "Do not attempt to answer from general knowledge if the information is not in the documents."
                )
                
                document_message = (
                    "DOCUMENT CONTENT START\n\n"
                    f"{document_content}\n\n"
                    "DOCUMENT CONTENT END\n\n"
                    "Please use ONLY the information between DOCUMENT CONTENT START and DOCUMENT CONTENT END "
                    "to answer my questions. If the information is not in the document, say so."
                )
                
                model_response1 = "I understand my role. I'll answer questions based strictly on the document content."
                model_response2 = "I've received the document content and will use it to answer your questions."
            
            # Start chat with document content and system instructions
            chat = model.start_chat(
                history=[
                    {
                        "role": "user",
                        "parts": [{"text": system_message}]
                    },
                    {
                        "role": "model",
                        "parts": [{"text": model_response1}]
                    },
                    {
                        "role": "user",
                        "parts": [{"text": document_message}]
                    },
                    {
                        "role": "model",
                        "parts": [{"text": model_response2}]
                    }
                ]
            )
            st.session_state["chat"] = chat
        
        # Initialize messages if not already done
        if "messages" not in st.session_state:
            st.session_state["messages"] = []

        # Display existing messages
        for message in st.session_state["messages"]:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Handle new user input with appropriate placeholder text
        chat_placeholder = "दस्तऐवजांबद्दल प्रश्न विचारा:" if language == "मराठी" else "Ask a question about the documents:"
        if prompt := st.chat_input(chat_placeholder):
            st.session_state["messages"].append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""
                try:
                    response = st.session_state["chat"].send_message(prompt, stream=True)
                    for chunk in response:
                        full_response += (chunk.text if chunk.text else "")
                        message_placeholder.markdown(full_response + "▌")
                    message_placeholder.markdown(full_response)
                    st.session_state["messages"].append({"role": "assistant", "content": full_response})
                except Exception as e:
                    error_msg = f"संदेश निर्मितीदरम्यान त्रुटी: {e}" if language == "मराठी" else f"Error during message generation: {e}"
                    st.error(error_msg)

    except Exception as e:
        error_msg = f"मॉडेल आरंभ करताना त्रुटी: {e}" if language == "मराठी" else f"Error initializing model: {e}"
        st.error(error_msg)

# ========== STEP 5: Streamlit App ==========

def main():
    # Language selection
    language = st.sidebar.radio(
        "भाषा निवडा / Select Language",
        ["मराठी", "English"]
    )
    
    if language == "मराठी":
        st.title("जेमिनी सह दस्तऐवज चॅटबॉट")
        st.subheader("प्रश्न विचारण्यासाठी PDF आणि DOCX फाइल्स अपलोड करा.")

        # Add a sidebar with information in Marathi
        with st.sidebar:
            st.header("माहिती")
            st.write("""
            हे अॅप्लिकेशन आपल्याला PDF आणि DOCX फाइल्स अपलोड करण्यास आणि त्यांच्या सामग्रीबद्दल प्रश्न विचारण्यास अनुमती देते.
            जेमिनी AI मॉडेल दस्तऐवजांचे विश्लेषण करेल आणि सामग्रीच्या आधारे आपल्या प्रश्नांची उत्तरे देईल.
            """)
            
            st.header("टिप्स")
            st.write("""
            - सर्वोत्तम परिणामांसाठी, स्पष्ट, सुव्यवस्थित मजकूरासह दस्तऐवज अपलोड करा
            - दस्तऐवज सामग्रीशी संबंधित विशिष्ट प्रश्न विचारा
            - जर आपल्याला अपूर्ण उत्तरे मिळत असतील, तर आपला प्रश्न पुन्हा मांडण्याचा प्रयत्न करा
            - मॉडेलला संदर्भ मर्यादा आहे, म्हणून खूप मोठे दस्तऐवज पूर्णपणे प्रक्रिया केले जाऊ शकत नाहीत
            """)
    else:
        st.title("Document Chatbot with Gemini")
        st.subheader("Upload PDF and DOCX files to ask questions.")

        # Add a sidebar with information in English
        with st.sidebar:
            st.header("About")
            st.write("""
            This application allows you to upload PDF and DOCX files and ask questions about their content.
            The Gemini AI model will analyze the documents and answer your questions based on the content.
            """)
            
            st.header("Tips")
            st.write("""
            - For best results, upload documents with clear, well-formatted text
            - Ask specific questions related to the document content
            - If you get incomplete answers, try rephrasing your question
            - The model has a context limit, so very large documents may not be fully processed
            """)

    # Store the selected language in session state to maintain it across reruns
    if "language" not in st.session_state:
        st.session_state["language"] = language
    else:
        if language != st.session_state["language"]:
            st.session_state["language"] = language
            # Clear chat when language changes
            if "messages" in st.session_state:
                st.session_state.pop("messages")
            if "chat" in st.session_state:
                st.session_state.pop("chat")
    
    # Use the appropriate language for UI elements
    if st.session_state["language"] == "मराठी":
        pdf_label = "PDF फाइल अपलोड करा"
        docx_label = "DOCX फाइल अपलोड करा"
        processing_msg = "दस्तऐवज प्रक्रिया करत आहे..."
        details_label = "दस्तऐवज प्रक्रिया तपशील"
        success_msg = "दस्तऐवज प्रक्रिया यशस्वीरित्या पूर्ण झाली."
        length_msg = "एकूण सामग्री लांबी: {} वर्ण"
        file_msg = "{} फाइल: {}"
        reset_btn = "चॅट रीसेट करा"
        reset_msg = "चॅट इतिहास रीसेट केला गेला आहे."
        error_msg = "दस्तऐवज प्रक्रिया करण्यात अयशस्वी. कृपया वैध फाइल्ससह पुन्हा प्रयत्न करा."
        upload_msg = "चॅटबॉट सुरू करण्यासाठी कृपया PDF किंवा DOCX फाइल्स अपलोड करा."
    else:
        pdf_label = "Upload a PDF file"
        docx_label = "Upload a DOCX file"
        processing_msg = "Processing documents..."
        details_label = "Document Processing Details"
        success_msg = "Document processing completed successfully."
        length_msg = "Total content length: {} characters"
        file_msg = "{} file: {}"
        reset_btn = "Reset Chat"
        reset_msg = "Chat history has been reset."
        error_msg = "Failed to process document content. Please try again with valid files."
        upload_msg = "Please upload PDF or DOCX files to start the chatbot."
    
    uploaded_pdf = st.file_uploader(pdf_label, type=["pdf"])
    uploaded_docx = st.file_uploader(docx_label, type=["docx"])

    if uploaded_pdf or uploaded_docx:
        with st.spinner(processing_msg):
            document_content = process_document_content(uploaded_pdf, uploaded_docx)
            if document_content:
                # Show a collapsible section with document details
                with st.expander(details_label):
                    st.write(success_msg)
                    st.write(length_msg.format(len(document_content)))
                    if uploaded_pdf:
                        st.write(file_msg.format("PDF", uploaded_pdf.name))
                    if uploaded_docx:
                        st.write(file_msg.format("DOCX", uploaded_docx.name))
                
                # Pass the selected language to the chatbot interface
                chatbot_interface(document_content, st.session_state["language"])

                # Optional: Add a button to reset the chat
                if st.button(reset_btn):
                    try:
                        # Clear session state
                        if "messages" in st.session_state:
                            st.session_state.pop("messages")
                        if "chat" in st.session_state:
                            st.session_state.pop("chat")
                        st.info(reset_msg)
                        # Force a rerun to clear the chat history
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.error(error_msg)
    else:
        st.info(upload_msg)

if __name__ == "__main__":
    main()